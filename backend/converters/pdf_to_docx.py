"""
PDF to Word.

The usual library for this (pdf2docx) depends on OpenCV, which is ~90 MB and
would eat most of the free tier's memory budget. This takes a lighter route:
PyMuPDF for text, tables and font metrics, python-docx to write the file.

The trade-off is honest — you get editable text with headings, paragraphs and
tables, not a pixel-perfect reproduction of the original layout.
"""

import io

import pymupdf

MAX_PAGES = 200

# A span must be this much larger than body text to count as a heading.
HEADING_RATIOS = ((1.6, 1), (1.35, 2), (1.15, 3))


def _body_font_size(document: pymupdf.Document) -> float:
    """
    Find the document's body text size.

    Counting how many *spans* use each size gets this wrong on short documents:
    a page with three headings and three paragraphs ties, and the tie can pick
    the heading size as "body", after which nothing looks like a heading.
    Weighting by character count reflects that body text is simply more of the
    document's text, however few paragraphs there are.
    """
    weight: dict[float, int] = {}
    for page in document:
        for block in page.get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span["text"].strip()
                    if text:
                        size = round(span["size"], 1)
                        weight[size] = weight.get(size, 0) + len(text)

    if not weight:
        return 11.0
    return max(weight.items(), key=lambda item: item[1])[0]


def _heading_level(size: float, body_size: float, bold: bool) -> int | None:
    ratio = size / body_size if body_size else 1
    for threshold, level in HEADING_RATIOS:
        if ratio >= threshold:
            return level
    # Same size but bold and short usually means a run-in heading.
    return 4 if bold and ratio >= 1.02 else None


def _blocks_outside_tables(page: pymupdf.Page, table_rects: list[pymupdf.Rect]):
    for block in page.get_text("dict")["blocks"]:
        if "lines" not in block:
            continue
        rect = pymupdf.Rect(block["bbox"])
        if any(rect.intersects(t) for t in table_rects):
            continue
        yield block


def pdf_to_docx(file_bytes: bytes) -> bytes:
    from docx import Document  # lazy import
    from docx.shared import Pt

    document = pymupdf.open(stream=file_bytes, filetype="pdf")
    try:
        if document.page_count == 0:
            raise ValueError("This PDF has no pages")
        if document.page_count > MAX_PAGES:
            raise ValueError(
                f"This PDF has {document.page_count} pages — the limit is {MAX_PAGES}"
            )

        body_size = _body_font_size(document)
        word = Document()
        wrote_anything = False

        for index, page in enumerate(document):
            tables = page.find_tables(strategy="lines").tables
            table_rects = [pymupdf.Rect(t.bbox) for t in tables]

            for block in _blocks_outside_tables(page, table_rects):
                for line in block["lines"]:
                    spans = [s for s in line["spans"] if s["text"].strip()]
                    if not spans:
                        continue

                    text = "".join(s["text"] for s in spans).strip()
                    if not text:
                        continue

                    largest = max(spans, key=lambda s: s["size"])
                    # Bit 4 of the span flags marks bold in PyMuPDF.
                    bold = bool(largest["flags"] & 2 ** 4)
                    level = _heading_level(largest["size"], body_size, bold)

                    if level:
                        word.add_heading(text, level=level)
                    else:
                        paragraph = word.add_paragraph(text)
                        paragraph.runs[0].font.size = Pt(round(largest["size"]))
                    wrote_anything = True

            for table in tables:
                rows = [
                    [("" if c is None else str(c)).strip() for c in row]
                    for row in table.extract()
                ]
                rows = [r for r in rows if any(r)]
                if len(rows) < 2:
                    continue

                word_table = word.add_table(rows=0, cols=len(rows[0]))
                word_table.style = "Table Grid"
                for row in rows:
                    cells = word_table.add_row().cells
                    for cell, value in zip(cells, row):
                        cell.text = value
                wrote_anything = True

            if index < document.page_count - 1:
                word.add_page_break()

        if not wrote_anything:
            raise ValueError(
                "No text found — this PDF may be scanned images, which needs OCR"
            )

        buffer = io.BytesIO()
        word.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    finally:
        document.close()
